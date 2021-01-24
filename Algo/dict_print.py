#GRABS ALL THE STANDARD DICTS, THEN COMPILES THEM SO HUMANS CAN SORT AND CLEAN THEM BY HAND
import re
import pandas as pd
from pandas import ExcelWriter
from pandas import ExcelFile
import psycopg2
import os
from itertools import groupby
from operator import itemgetter
from docx.enum.text import WD_COLOR_INDEX
from docx import Document

connection = psycopg2.connect(host="ec2-34-197-188-147.compute-1.amazonaws.com", dbname="d7p3fuehaleleo", user="snbetggfklcniv", password="7798f45239eda70f8278ce3c05dc632ad57b97957b601681a3c516f37153403a")
connection.autocommit = True
cursor = connection.cursor()

scrape_query = "select * from old_stand_dict;"
cursor.execute(scrape_query)
entries = cursor.fetchall()

#BALANCE SHEET
balance = ['Cash And Equivalents',
'Restricted Cash',
'Short Term Investments',
'Trading Asset Securities',
'Total Cash & ST Investments',
'Accounts Receivable',
'Notes Receivable',
'Other Receivables',
'Total Receivables',
'Inventory',
'Deferred Tax Assets, Curr.',
'Other Current Assets',
'Total Current Assets',
'Gross Property, Plant & Equipment',
'Accumulated Depreciation',
'Net Property, Plant & Equipment',
'Long-term Investments',
'Goodwill',
'Other Intangibles',
'Deferred Tax Assets, LT',
'Other Long-Term Assets',
'Total Assets','Accounts Payable',
'Accrued Exp.',
'Short-term Borrowings',
'Curr. Port. of LT Debt',
'Curr. Income Taxes Payable',
'Unearned Revenue, Current',
'Interest Capitalized',
'Other Current Liabilities',
'Total Current Liabilities',
'Long-Term Debt',
'Capital Leases',
'Unearned Revenue, Non-Current',
'Def. Tax Liability, Non-Curr.',
'Other Non-Current Liabilities',
'Total Liabilities',
'Common Stock',
'Additional Paid In Capital',
'Retained Earnings',
'Treasury Stock',
'Comprehensive Inc. and Other',
'Total Common Equity',
'Total Shares Out.',
'Total Equity',
'Total Liabilities And Equity']

#INCOME STATEMENT
income = [
'Revenue',
'Other Revenue',
'Total Revenue',
'Cost Of Goods Sold',
'Gross Profit',
'Selling General & Admin Exp. ',
'R & D Exp.',
'Depreciation & Amort.',
'Other Operating Expense/(Income)',
'Operating Exp., Total',
'Operating Income',
'Interest Expense',
'Interest and Invest. Income',
'Net Interest Exp.',
'Currency Exchange Gains',
'Other Non-Operating Inc. (Exp.)',
'EBT Excl. Unusual Items',
'Restructuring Charges',
'Merger & Related Restruct. Charges',
'Impairment of Goodwill',
'Gain (Loss) On Sale Of Invest.',
'Asset Writedown',
'Other Unusual Items',
'EBT Incl. Unusual Items',
'Income Tax Expense',
'Earnings from Cont. Ops.',
'Earnings of Discontinued Ops.',
'Other Changes',
'Net Income to Company',
'Minority Int. in Earnings',
'Net Income',
'Pref. Dividends and Other Adj.','Basic EPS',
'Weighted Avg. Basic Shares Out.',
'Diluted EPS',
'Weighted Avg. Diluted Shares Out.']

#CASH FLOW

cash = [
'Net Income',
'Depreciation & Amort.',
'Amort. of Goodwill and Intangibles',
'Depreciation & Amort., Total',
'Stock-Based Compensation',
'Net Cash From Discontinued Ops.',
'Other Operating Activities',
'Change in Acc. Receivable',
'Change In Inventories',
'Change in Acc. Payable',
'Change in Unearned Rev.',
'Cash from Ops.',
'Capital Expenditure',
'Cash Acquisitions',
'Divestitures',
'Purchase of Marketable Equity Securities',
'Sale of Marketable Equity Securities',
'Net (Inc.) Dec. in Loans Originated/Sold',
'Other Investing Activities',
'Cash from Investing',
'Short Term Debt Issued',
'Long-Term Debt Issued',
'Total Debt Issued',
'Short Term Debt Repaid',
'Long-Term Debt Repaid',
'Total Debt Repaid',
'Repurchase of Common Stock',
'Issuance of Common Stock',
'Common Dividends Paid',
'Total Dividends Paid',
'Special Dividend Paid',
'Other Financing Activities',
'Cash from Financing',
'Foreign Exchange Rate Adj.',
'Net Change in Cash']


#document = Document()

copy = []
for entry in entries:
    dict = {}
    dict['standard_name'] = entry[0]
    dict['acc_name'] = entry[1]
    copy.append(dict)

seen = set()
new_l = []
for d in copy:
    t = tuple(d.items())
    if t not in seen:
        seen.add(t)
        new_l.append(d)

s = 0
k = itemgetter('standard_name')
i = groupby(sorted(new_l, key=k), key=k)

finale = []
for key, new_l in i:
    s+=1
    header = str(key)
    print(header)
    b_check = [item for item in balance if item == header]
    c_check = [item for item in cash if item == header]
    i_check = [item for item in income if item == header]
    if i_check:
        print('INCOME')
    if c_check:
        print('CASH')
    if b_check:
        print('BALANCE')
    #document.add_paragraph(header)
    for item in new_l:
        if i_check:
            copy = item.copy()
            copy['statement'] = 'income'
            print(copy)
            finale.append(copy)
            insert_state = 'standard_dict'
            standard_name = copy['standard_name']
            acc_name = copy['acc_name']
            statement = copy['statement']
            sql_statement = "INSERT INTO %s (standard_name, acc_name, statement) VALUES('%s', '%s', '%s')"%(insert_state,standard_name, acc_name, statement)
            print(sql_statement)
            cursor.execute(sql_statement)
        if c_check:
            copy = item.copy()
            copy['statement'] = 'cash_flow'
            print(copy)
            finale.append(copy)
            insert_state = 'standard_dict'
            standard_name = copy['standard_name']
            acc_name = copy['acc_name']
            statement = copy['statement']
            sql_statement = "INSERT INTO %s (standard_name, acc_name, statement) VALUES('%s', '%s', '%s')"%(insert_state,standard_name, acc_name, statement)
            print(sql_statement)
            cursor.execute(sql_statement)
        if b_check:
            copy = item.copy()
            copy['statement'] = 'balance'
            print(copy)
            finale.append(copy)
            insert_state = 'standard_dict'
            standard_name = copy['standard_name']
            acc_name = copy['acc_name']
            statement = copy['statement']
            sql_statement = "INSERT INTO %s (standard_name, acc_name, statement) VALUES('%s', '%s', '%s')"%(insert_state,standard_name, acc_name, statement)
            print(sql_statement)
            cursor.execute(sql_statement)
        #document.add_paragraph(smith)

#print(s)
#document.save('docx_file.docx')
